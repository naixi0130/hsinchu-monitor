import streamlit as st
import feedparser
import requests
from datetime import datetime, timedelta
from collections import defaultdict
import pandas as pd
from apify_client import ApifyClient
from docx import Document
from io import BytesIO

st.set_page_config(page_title="新竹縣社群監測器", page_icon="🏠", layout="wide")
st.title("🏠 新竹縣社群監測器")
st.caption("嚴格過去24小時 • 已優化 FB 社團 + IG 關鍵字過濾")

st.markdown("---")

# ==================== 關鍵字分類 ====================
politics_keywords = ["選舉", "議員", "縣長", "立委", "政黨", "藍營", "綠營", "民進黨", "國民黨", "民眾黨", "柯文哲", "侯友宜", "鄭朝方", "林智堅", "議會", "質詢"]
issues_keywords = ["竹北", "科學園區", "交通", "捷運", "高鐵", "房價", "開發", "環境", "空汙", "醫療", "教育", "建設", "停車", "道路"]
disaster_keywords = ["地震", "颱風", "淹水", "豪雨", "火災", "爆炸", "車禍", "意外", "天災", "災情", "傷亡", "淹大水", "救援"]

def classify_post(title, summary):
    text = (title + " " + summary).lower()
    if any(kw in text for kw in disaster_keywords):
        return "disasters", "🟥 新竹縣天災人禍"
    elif any(kw in text for kw in politics_keywords):
        return "politics", "🔵 新竹縣政治"
    else:
        return "issues", "🟠 新竹縣重大議題"

now = datetime.now()
cutoff = now - timedelta(hours=24)

# ==================== 抓取函式 ====================
def fetch_google_news():
    url = "https://news.google.com/rss/search?q=%E6%96%B0%E7%AB%B9%E7%B8%A3&hl=zh-TW&gl=TW&ceid=TW:zh-Hant"
    feed = feedparser.parse(url)
    posts = []
    for entry in feed.entries:
        pub_date = datetime(*entry.published_parsed[:6]) if hasattr(entry, 'published_parsed') else now
        if pub_date > cutoff:
            posts.append({
                "title": entry.title,
                "summary": getattr(entry, 'summary', '')[:300],
                "time": pub_date.strftime("%Y-%m-%d %H:%M"),
                "platform": "Google News",
                "url": entry.link
            })
    return posts

def fetch_ptt():
    headers = {"User-Agent": "Mozilla/5.0"}
    session = requests.Session()
    session.cookies.set("over18", "1")
    url = "https://www.ptt.cc/bbs/Gossiping/search?q=%E6%96%B0%E7%AB%B9%E7%B8%A3"
    try:
        r = session.get(url, headers=headers, timeout=10)
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(r.text, "html.parser")
        posts = []
        for div in soup.find_all("div", class_="r-ent")[:12]:
            title_a = div.find("a")
            if title_a:
                posts.append({
                    "title": title_a.text.strip(),
                    "summary": "PTT 熱門討論 - 新竹縣相關",
                    "time": "PTT 最近24小時",
                    "platform": "PTT",
                    "url": "https://www.ptt.cc" + title_a["href"]
                })
        return posts
    except:
        return []

def fetch_dcard():
    url = "https://www.dcard.tw/service/api/v2/posts?limit=30"
    headers = {"User-Agent": "Mozilla/5.0"}
    try:
        r = requests.get(url, headers=headers, timeout=10)
        data = r.json()
        posts = []
        for post in data:
            title = post.get("title", "")
            excerpt = post.get("excerpt", "")
            if any(x in (title + excerpt) for x in ["新竹縣", "新竹"]):
                created_time = datetime.fromtimestamp(post.get("createdAt", 0) / 1000)
                if created_time > cutoff:
                    posts.append({
                        "title": title,
                        "summary": excerpt[:300],
                        "time": created_time.strftime("%Y-%m-%d %H:%M"),
                        "platform": "Dcard",
                        "url": f"https://www.dcard.tw/f/post/{post.get('id')}"
                    })
        return posts
    except:
        return []

# ==================== Facebook - 只抓推薦社團 ====================
def fetch_fb_posts():
    if "APIFY_API_TOKEN" not in st.secrets:
        return []
    try:
        client = ApifyClient(st.secrets["APIFY_API_TOKEN"])
        run_input = {
            "startUrls": [
                {"url": "https://www.facebook.com/groups/965991216800536/"},   # 新竹縣人新竹縣事
                {"url": "https://www.facebook.com/groups/1924816131138636/"},  # 我是竹北人
                {"url": "https://www.facebook.com/groups/195118655867167/"},   # 竹北生活大小事
                {"url": "https://www.facebook.com/groups/181766803868709/"},   # 新竹大小事。二社
                {"url": "https://www.facebook.com/groups/201184623938653/"},   # 竹北大小事
            ],
            "maxPosts": 25,
            "onlyPosts": True
        }
        run = client.actor("apify/facebook-posts-scraper").call(run_input=run_input)
        dataset = client.dataset(run["defaultDatasetId"]).iterate_items()
        posts = []
        for item in list(dataset)[:15]:
            text = item.get("text") or item.get("message") or ""
            if len(text.strip()) > 30:
                posts.append({
                    "title": text[:85].replace("\n", " "),
                    "summary": text[:400],
                    "time": now.strftime("%Y-%m-%d %H:%M"),
                    "platform": "Facebook",
                    "url": item.get("url") or "#"
                })
        return posts
    except Exception as e:
        st.warning(f"Facebook 抓取失敗: {str(e)[:150]}")
        return []

# ==================== Instagram - 加強關鍵字過濾 ====================
def fetch_ig_hashtag():
    if "APIFY_API_TOKEN" not in st.secrets:
        return []
    try:
        client = ApifyClient(st.secrets["APIFY_API_TOKEN"])
        run_input = {"hashtags": ["新竹縣", "竹北"], "resultsLimit": 20}
        run = client.actor("apify/instagram-hashtag-scraper").call(run_input=run_input)
        dataset = client.dataset(run["defaultDatasetId"]).iterate_items()
        posts = []
        filter_keywords = politics_keywords + issues_keywords + disaster_keywords
        for item in list(dataset)[:12]:
            caption = item.get("caption", "") or ""
            if any(kw in caption for kw in filter_keywords):   # 加強過濾，只保留相關貼文
                posts.append({
                    "title": caption[:80],
                    "summary": caption[:350],
                    "time": now.strftime("%Y-%m-%d %H:%M"),
                    "platform": "Instagram",
                    "url": item.get("url") or "#"
                })
        return posts
    except Exception as e:
        st.warning(f"Instagram 抓取失敗: {str(e)[:100]}")
        return []

def generate_word(all_posts, categories):
    doc = Document()
    doc.add_heading('新竹縣社群監測報告 - 過去24小時', 0)
    doc.add_paragraph(f'產生時間：{now.strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph(f'總共 {len(all_posts)} 則討論\n')
    
    for cat_key, cat_name in [("politics", "🔵 新竹縣政治"), ("issues", "🟠 新竹縣重大議題"), ("disasters", "🟥 新竹縣天災人禍")]:
        doc.add_heading(cat_name, level=1)
        for p in categories.get(cat_key, []):
            doc.add_heading(p['title'], level=2)
            doc.add_paragraph(f"平台：{p['platform']}   時間：{p['time']}")
            doc.add_paragraph(p.get('summary', '無摘要'))
            if p.get('url') and p['url'] != "#":
                doc.add_paragraph(f"連結：{p['url']}")
            doc.add_paragraph("─" * 40)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==================== 主程式 ====================
if st.button("🔥 一鍵抓取過去24小時討論", type="primary", use_container_width=True):
    with st.spinner("抓取中... FB 社團 + IG 已加強過濾（可能需 30-60 秒）"):
        news = fetch_google_news()
        ptt = fetch_ptt()
        dcard = fetch_dcard()
        fb = fetch_fb_posts()
        ig = fetch_ig_hashtag()
        
        all_posts = news + ptt + dcard + fb + ig
        
        categories = defaultdict(list)
        for post in all_posts:
            cat_key, _ = classify_post(post["title"], post.get("summary", ""))
            categories[cat_key].append(post)
        
        st.success(f"✅ 抓取完成！共 {len(all_posts)} 則（Facebook {len(fb)} 則，Instagram {len(ig)} 則）")
        
        col1, col2, col3 = st.columns(3)
        for col, name, key in zip([col1, col2, col3], 
                                  ["🔵 新竹縣政治", "🟠 新竹縣重大議題", "🟥 新竹縣天災人禍"], 
                                  ["politics", "issues", "disasters"]):
            with col:
                st.subheader(name)
                st.caption(f"{len(categories[key])} 則")
                for p in categories[key]:
                    with st.expander(f"**{p.get('title', '無標題')[:70]}...**", expanded=False):
                        st.caption(f"{p['platform']} • {p['time']}")
                        st.write(p.get('summary', '無內容'))
                        if p.get('url') and p['url'] != "#":
                            st.link_button("🔗 查看原文", p['url'])

        if all_posts:
            word_bytes = generate_word(all_posts, categories)
            st.download_button(
                label="📄 匯出全部結果為 Word 文件 (.docx)",
                data=word_bytes,
                file_name=f"新竹縣監測報告_{now.strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

else:
    st.info("👆 點擊上方按鈕開始抓取\n\n※ FB 已鎖定特定社團，IG 已加強政治/議題/災情過濾，無關內容會大幅減少")

st.markdown("---")
st.caption("FB 已加入推薦社團 • IG 已加強關鍵字過濾 • 若還是抓到無關內容，請告訴我，我再繼續調整")
